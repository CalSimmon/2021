# Imports
## Besides default libraries, you'll need to install three:
### 1. PyPDF2
### 2. pdf2image
### 3. python-docx
## All can be installed using pip

import os
from os import listdir
from os.path import isfile, join
from pathlib import Path
from PyPDF2 import PdfFileWriter, PdfFileReader
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from datetime import datetime

# Check to make sure the PDFs folder exists.
# Error if not
if not os.path.exists(".\PDFs"):
    print("PDFs folder does not exist.  Please create a folder called 'PDFs and place that in the containing folder.'")
    quit()

# Continue if so
else:
    print("PDFs folder exists")

# Check if CroppedImages folder exists
# Create folder if not
if not os.path.exists(".\CroppedImages"):
    print("Creating CroppedImages Folder")
    os.mkdir("CroppedImages")

# Continue if so
else:
    print("CroppedImages folder exists")

# Establish path of PDFs and get all files contained within
path = ".\PDFs"
onlyfiles = [f for f in listdir(path) if isfile(join(path, f))]

# Run through each PDF and process
for file in onlyfiles:
    print(file)
    pathfile = ".\PDFs\\" + file  # Create path for each file

    pdf_reader = PdfFileReader(str(pathfile), strict=False)
    numberPages = pdf_reader.numPages  # Enter into PDF Reader and get page number

    for x in range(numberPages):  # Run through each page and process
        first_page = pdf_reader.getPage(x)

        # Select the section to crop
        first_page.mediaBox.upperLeft = (0, 480)
        first_page.mediaBox.lowerRight = (150, 400)

        # Create the filename and new path in the CroppedImages folder
        fileExt = "_Cropped_Page" + str(x + 1) + ".pdf"
        croppedfilename = file.replace('.pdf', fileExt)
        outpath = ".\CroppedImages\\" + croppedfilename

        # Write the file to a new PDF
        pdf_writer = PdfFileWriter()
        pdf_writer.addPage(first_page)
        with Path(outpath).open(mode="wb") as output_file:
            pdf_writer.write(output_file)

        # Use convert_from_path to convert the PDF to a JPEG
        images = convert_from_path(outpath)
        for i in range(len(images)):
            images[i].save(outpath.replace('.pdf', '.jpg'))

        # Remove the cropped PDF to leave the JPEG
        if os.path.exists(outpath):
            os.remove(outpath)
        else:
            print("This file does not exist")


# Use Document() to create a new document to consolidate all images
document = Document()

# Establish the path to images, and add all image names into a variable
path = ".\CroppedImages"
croppedImageFiles = [f for f in listdir(path) if isfile(join(path, f))]

# Run through all files within the CroppedImages folder and process
for cropimages in croppedImageFiles:
    document.add_paragraph(cropimages)  # Add filename labels to each image
    imgPath = ".\CroppedImages\\" + cropimages
    document.add_picture(imgPath, width=Inches(5))  # Add image to document

# Use datetime to create the document name
dateTimeObj = datetime.now()
documentName = dateTimeObj.strftime("%Y%m%d") + "_PDFImageDump.docx"
document.save(documentName)  # Save document using establish name
